"""Microsoft Word (.docx) parser."""

from __future__ import annotations

import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path

from docx import Document

from doc_analyzer.models import DocumentMetrics, Heading
from doc_analyzer.utils.text_utils import (
    detect_hierarchy_issues,
    heading_level_from_style,
    sentence_stats,
    word_count,
)

_NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}


def parse_docx(path: Path) -> DocumentMetrics:
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(f"DOCX not found: {path}")
    if path.stat().st_size == 0:
        raise ValueError(f"Empty file: {path}")

    doc = Document(str(path))
    paragraphs_text: list[str] = []
    headings: list[Heading] = []
    bullet_count = 0
    links: list[str] = []
    table_count = 0

    total_paragraphs = len(doc.paragraphs)
    empty_count = 0

    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if not text:
            empty_count += 1
            continue
        paragraphs_text.append(text)
        style = para.style.name if para.style else ""
        level = heading_level_from_style(style)
        if level:
            headings.append(Heading(level=level, text=text, style=style))
        if style and "List" in style:
            bullet_count += 1

    for table in doc.tables:
        if any(cell.text.strip() for row in table.rows for cell in row.cells):
            table_count += 1

    links.extend(_extract_hyperlinks(path))
    links = list(dict.fromkeys(links))

    full_text = "\n".join(paragraphs_text)
    wc = word_count(full_text)
    sc, avg_sent = sentence_stats(full_text)
    para_count = len(paragraphs_text)
    avg_wpp = round(wc / para_count, 2) if para_count else 0.0
    empty_ratio = round(empty_count / total_paragraphs, 3) if total_paragraphs else 0.0
    levels = [h.level for h in headings]

    return DocumentMetrics(
        file_name=path.name,
        file_path=str(path.resolve()),
        file_type="DOCX",
        size_bytes=path.stat().st_size,
        page_count=None,
        word_count=wc,
        paragraph_count=para_count,
        heading_count=len(headings),
        table_count=table_count,
        image_count=0,
        link_count=len(links),
        bullet_count=bullet_count,
        sentence_count=sc,
        avg_words_per_paragraph=avg_wpp,
        avg_sentence_length=round(avg_sent, 2),
        empty_paragraph_ratio=empty_ratio,
        metadata=_core_properties(doc),
        headings=headings,
        heading_hierarchy_issues=detect_hierarchy_issues(levels),
        text_preview=full_text[:6000],
    )


def _core_properties(doc: Document) -> dict[str, str]:
    try:
        cp = doc.core_properties
        return {
            k: v
            for k, v in {
                "title": cp.title or "",
                "author": cp.author or "",
                "subject": cp.subject or "",
                "created": str(cp.created) if cp.created else "",
                "modified": str(cp.modified) if cp.modified else "",
                "keywords": cp.keywords or "",
            }.items()
            if v
        }
    except Exception:
        return {}


def _extract_hyperlinks(path: Path) -> list[str]:
    found: list[str] = []
    try:
        with zipfile.ZipFile(path) as archive:
            if "word/document.xml" not in archive.namelist():
                return found
            root = ET.fromstring(archive.read("word/document.xml"))
            for hyperlink in root.findall(".//w:hyperlink", _NS):
                for text_node in hyperlink.findall(".//w:t", _NS):
                    if text_node.text:
                        found.append(text_node.text)
    except Exception:
        pass
    return found
